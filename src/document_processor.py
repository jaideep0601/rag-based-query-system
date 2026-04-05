import os
from typing import List


class Document:
    def __init__(self, page_content: str, metadata: dict):
        self.page_content = page_content
        self.metadata = metadata


class DocumentProcessor:
    
    def __init__(self, chunk_size=1000, chunk_overlap=200):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_text(self, text: str, source: str) -> List[Document]:
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunk_text = text[start:end]
            
            if chunk_text.strip():
                chunks.append(Document(
                    page_content=chunk_text,
                    metadata={"source": source}
                ))
            
            start += self.chunk_size - self.chunk_overlap
        
        print(f"Split into {len(chunks)} chunks")
        return chunks
    
    def load_txt(self, file_path: str) -> List[Document]:
        print(f"Loading TXT: {file_path}")
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        return self.split_text(text, file_path)
    
    def load_pdf(self, file_path: str) -> List[Document]:
        print(f"Loading PDF: {file_path}")
        try:
            import pypdf
            reader = pypdf.PdfReader(file_path)
            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"
            return self.split_text(text, file_path)
        except ImportError:
            raise ImportError("Install pypdf: pip install pypdf")
    
    def load_docx(self, file_path: str) -> List[Document]:
        print(f"Loading DOCX: {file_path}")
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = '\n'.join([para.text for para in doc.paragraphs])
            return self.split_text(text, file_path)
        except ImportError:
            raise ImportError("Install python-docx: pip install python-docx")
    
    def load_document(self, file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        extension = file_path.lower().split('.')[-1]
        
        if extension == 'pdf':
            return self.load_pdf(file_path)
        elif extension == 'txt':
            return self.load_txt(file_path)
        elif extension == 'docx':
            return self.load_docx(file_path)
        else:
            raise ValueError(f"Unsupported file type: {extension}")